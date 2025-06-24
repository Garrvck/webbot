from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches

def generate_resume_doc(resume: dict):
    def hide_borders(table):
        for row in table.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'nil')
                    tcBorders.append(border)
                tcPr.append(tcBorders)

    def set_font(paragraph, size=12, bold=False):
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(size)
            run.bold = bold
        paragraph.style.font.name = 'Times New Roman'

    def add_field(paragraph, label, value):
        run_bold = paragraph.add_run(label)
        run_bold.bold = True
        run_bold.font.name = 'Times New Roman'
        run_bold.font.size = Pt(12)

        run_normal = paragraph.add_run(" " + str(value))
        run_normal.font.name = 'Times New Roman'
        run_normal.font.size = Pt(12)

    doc = Document()

    # Margins
    section = doc.sections[0]
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2)

    # 1-jadval: sarlavha, ism, lavozim, rasm
    table = doc.add_table(rows=3, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(13.5)
    table.columns[1].width = Cm(3.5)

    texts = [
        ("MA’LUMOTNOMA", True),
        (resume["full_name"], True),
        (resume.get("specialization", ""), False)
    ]

    for i in range(3):
        p = table.cell(i, 0).paragraphs[0]
        run = p.add_run(texts[i][0])
        run.bold = texts[i][1]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    merged_cell = table.cell(0, 1).merge(table.cell(2, 1))
    photo_path = resume.get("photo_path")
    if photo_path and photo_path != "Yuklanmagan":
        try:
            merged_cell.paragraphs[0].clear()
            merged_cell.paragraphs[0].add_run().add_picture(photo_path, width=Inches(1.2))
        except:
            merged_cell.paragraphs[0].add_run("[Rasmni ochib bo‘lmadi]")
    else:
        merged_cell.paragraphs[0].add_run("[Rasm]")




    # merged_cell = table.cell(0, 1).merge(table.cell(2, 1))
    # merged_cell.paragraphs[0].add_run("[Rasm]")  # Agar rasm yuklanmasa

    doc.add_paragraph("")
    doc.add_paragraph("")

    # 2-jadval: foydalanuvchi ma'lumotlari
    table = doc.add_table(rows=6, cols=2)
    hide_borders(table)

    data = [
        [("Tug‘ilgan yili:", resume.get("birth_date", "")), ("Tug‘ilgan joyi:", resume.get("birth_place", ""))],
        [("Millati:", resume.get("nationality", "")), ("Partiyaviyligi:", resume.get("party_membership", ""))],
        [("Ma’lumoti:", resume.get("education", "")), ("Tamomlagan:", resume.get("university", ""))],
        [("Ma’lumoti bo‘yicha mutaxassisligi:", resume.get("specialization", "")), ("Ilmiy daraja:", resume.get("ilmiy_daraja", ""))],
        [("Ilmiy unvon:", resume.get("ilmiy_unvon", "")), ("Qaysi chet tillarini biladi:", resume.get("languages", ""))],
        [("Saylanadigan organlar a’zoligi:", resume.get("deputat", "")), ("Davlat mukofotlari:", resume.get("dav_mukofoti", ""))]
    ]

    for row_idx, row_data in enumerate(data):
        for col_idx, (label, value) in enumerate(row_data):
            p = table.cell(row_idx, col_idx).paragraphs[0]
            add_field(p, label, value)

    doc.add_paragraph("")

    # Mehnat faoliyati
    p = doc.add_paragraph("MEHNAT FAOLIYATI")
    set_font(p, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    work_experience = resume.get("work_experience", "")
    p = doc.add_paragraph(work_experience)
    p.paragraph_format.left_indent = Pt(10)
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    doc.add_paragraph("")

    # Footerda telefon raqami
    footer = doc.sections[0].footer
    paragraph = footer.paragraphs[0]
    contact = resume.get("phone", "")
    paragraph.text = f"Tel: {contact}"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in paragraph.runs:
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    doc.add_paragraph("")
    doc.add_page_break()

    # Yaqin qarindoshlar
    p = doc.add_paragraph("Yaqin qarindoshlari haqida ma'lumot")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    relatives = resume.get("relatives", [])
    table = doc.add_table(rows=len(relatives) + 1, cols=5)
    table.style = 'Table Grid'

    headers = ["Yaqinlik darajasi", "F.I.Sh", "Tug‘ilgan sanasi", "Ish joyi va lavozimi", "Yashash manzili"]
    col_widths = [Cm(2.2), Cm(6.3), Cm(4), Cm(5), Cm(5)]

    for idx, (header, width) in enumerate(zip(headers, col_widths)):
        cell = table.cell(0, idx)
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    for row_idx, rel in enumerate(relatives, start=1):
        row_data = [
            rel.get("relation_type", ""),
            rel.get("full_name", ""),
            rel.get("b_year_place", ""),
            rel.get("job_title", ""),
            rel.get("address", "")
        ]
        for col_idx, value in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            run = cell.paragraphs[0].add_run(str(value))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

    filename = f"{resume.get('full_name', 'rezyume').replace(' ', '_')}_Malumotnoma.docx"
    doc.save(filename)
    print(f"✅ Word fayl yaratildi: {filename}")
    return filename
